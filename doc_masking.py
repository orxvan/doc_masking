import os
import shutil
import subprocess
import sys

from docx import Document


def convert_doc_with_libreoffice(folder_path, soffice_path):
    # (此函数与之前相同，无需修改)
    print("--- 阶段一: 开始使用 LibreOffice 将 .doc 转换为 .docx ---")
    soffice_cmd = soffice_path if soffice_path else "soffice"
    try:
        subprocess.run([soffice_cmd, "--version"], capture_output=True, check=True, text=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        print(f"[错误] 无法执行 '{soffice_cmd}' 命令。请检查LibreOffice安装和路径设置。")
        return False

    doc_files = [f for f in os.listdir(folder_path) if f.lower().endswith(".doc")]
    if not doc_files:
        print("未找到需要转换的 .doc 文件。")
        return True

    print(f"找到 {len(doc_files)} 个 .doc 文件需要转换。")
    for filename in doc_files:
        print(f"  正在转换: '{filename}'...")
        try:
            subprocess.run(
                [
                    soffice_cmd,
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    folder_path,
                    os.path.join(folder_path, filename),
                ],
                capture_output=True,
                check=True,
            )
            print("  成功转换为 .docx 版本。")
        except subprocess.CalledProcessError as e:
            print(f"  [错误] 转换 '{filename}' 失败: {e.stderr.decode('utf-8', 'ignore')}")
            continue
        except Exception as e:
            print(f"  [严重错误] 执行转换命令时发生意外: {e}")
            continue

    print("\n  正在清理原始 .doc 文件...")
    for filename in doc_files:
        doc_path = os.path.join(folder_path, filename)
        docx_path = os.path.join(folder_path, os.path.splitext(filename)[0] + ".docx")
        if os.path.exists(docx_path):
            try:
                os.remove(doc_path)
                print(f"  已删除: '{filename}'")
            except Exception as e:
                print(f"  [错误] 删除 '{filename}' 失败: {e}")

    print("--- 转换阶段完成 ---\n")
    return True


def replace_text_in_paragraph(p, keywords, replacement):
    """
    【核心修正函数】
    在一个段落(p)中查找并替换所有关键字。
    此函数可以处理关键字被分割到多个run中的情况。
    """
    # 拼接一个完整的文本字符串用于检测
    full_text = "".join(run.text for run in p.runs)

    found_keyword = False
    for keyword in keywords:
        if keyword in full_text:
            found_keyword = True
            # 在完整的文本上执行替换
            full_text = full_text.replace(keyword, replacement)

    # 如果检测并替换了任何关键字，就用新文本重构整个段落
    if found_keyword:
        # 保留原始段落的第一个run的格式
        original_style = p.runs[0].font if p.runs else None

        # 清空段落中所有的旧run
        p.clear()

        # 添加一个包含所有新文本的新run
        new_run = p.add_run(full_text)

        # 尝试将旧格式应用到新run上
        if original_style:
            new_run.font.name = original_style.name
            new_run.font.size = original_style.size
            new_run.bold = original_style.bold
            new_run.italic = original_style.italic
            new_run.underline = original_style.underline
            if original_style.color.rgb:
                new_run.font.color.rgb = original_style.color.rgb


def desensitize_docx_folder(folder_path, keywords, replacement):
    """
    (此函数已更新，调用新的修正函数)
    对文件夹内所有.docx文件进行文件名和内容的批量脱敏。
    """
    print("--- 阶段二: 开始对所有 .docx 文件进行脱敏 ---")

    # 步骤 2.1: 批量重命名文件 (逻辑不变)
    # ... (此处省略未改变的代码)
    print("  步骤 2.1: 正在重命名文件...")
    try:
        files_in_folder = [f for f in os.listdir(folder_path) if f.lower().endswith(".docx")]
        for filename in files_in_folder:
            new_filename = filename
            for keyword in keywords:
                new_filename = new_filename.replace(keyword, replacement)

            if new_filename != filename:
                original_filepath = os.path.join(folder_path, filename)
                new_filepath = os.path.join(folder_path, new_filename)
                shutil.move(original_filepath, new_filepath)
                print(f"    已重命名: '{filename}' -> '{new_filename}'")
    except Exception as e:
        print(f"  [错误] 文件重命名过程中发生错误: {e}")
    print("  文件重命名完成。\n")

    # 步骤 2.2: 批量替换文件内容 (已更新为更健壮的逻辑)
    print("  步骤 2.2: 正在替换文件内容...")
    docx_files = [f for f in os.listdir(folder_path) if f.lower().endswith(".docx")]
    for filename in docx_files:
        filepath = os.path.join(folder_path, filename)
        print(f"    正在处理内容: {filename}")
        try:
            doc = Document(filepath)

            # 处理所有普通段落
            for p in doc.paragraphs:
                replace_text_in_paragraph(p, keywords, replacement)

            # 处理所有表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_text_in_paragraph(p, keywords, replacement)

            doc.save(filepath)
        except Exception as e:
            print(f"    [错误] 处理 '{filename}' 内容时失败: {e}")

    print("--- 脱敏阶段完成 ---")


folder_path = r"./docs"


keywords_to_replace = ["特朗普", "川普"]

# 3. 设置替换后的文本
replacement_text = "XX公司"

# --- 修改结束 ---


if __name__ == "__main__":
    print("启动跨平台文档脱敏程序...")
    print("=" * 40)

    if not os.path.isdir(folder_path):
        print(f"[致命错误] 文件夹 '{folder_path}' 不存在，程序已中止。")
    else:
        # 阶段一：转换 .doc 文件
        conversion_successful = convert_doc_with_libreoffice(folder_path, soffice_path=None)

        # 只有在转换步骤成功或无需转换时才进行下一步
        if conversion_successful:
            # 阶段二：处理所有 .docx 文件
            desensitize_docx_folder(folder_path, keywords_to_replace, replacement_text)
            print("\n" + "=" * 40)
            print("所有任务已完成！")
        else:
            print("\n" + "=" * 40)
            print("由于转换阶段失败，脱敏程序已中止。")
